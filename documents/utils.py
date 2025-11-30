"""
Utilities for document import/export and conversion
"""
import re
from html.parser import HTMLParser
from django.utils.html import strip_tags
import markdown
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class HTMLToDocxConverter:
    """Convert HTML to DOCX"""
    
    @staticmethod
    def convert(html_content):
        """Convert HTML content to DOCX document"""
        doc = DocxDocument()
        
        # Simple HTML parsing - extract paragraphs and basic formatting
        # In production, use a proper HTML parser
        soup = _parse_html_simple(html_content)
        
        for element in soup:
            if element['type'] == 'p':
                para = doc.add_paragraph()
                for run_data in element.get('runs', []):
                    run = para.add_run(run_data['text'])
                    if run_data.get('bold'):
                        run.bold = True
                    if run_data.get('italic'):
                        run.italic = True
                    if run_data.get('underline'):
                        run.underline = True
            
            elif element['type'] == 'h1':
                para = doc.add_paragraph(element['text'], style='Heading 1')
            elif element['type'] == 'h2':
                para = doc.add_paragraph(element['text'], style='Heading 2')
            elif element['type'] == 'h3':
                para = doc.add_paragraph(element['text'], style='Heading 3')
            
            elif element['type'] == 'ul':
                for item in element.get('items', []):
                    para = doc.add_paragraph(item, style='List Bullet')
            
            elif element['type'] == 'ol':
                for i, item in enumerate(element.get('items', []), 1):
                    para = doc.add_paragraph(item, style='List Number')
        
        return doc


def _parse_html_simple(html_content):
    """Simple HTML parser - returns list of elements"""
    elements = []
    
    # Remove script and style tags
    html_content = re.sub(r'<script.*?</script>', '', html_content, flags=re.DOTALL)
    html_content = re.sub(r'<style.*?</style>', '', html_content, flags=re.DOTALL)
    
    # Extract body content if available
    body_match = re.search(r'<body[^>]*>(.*?)</body>', html_content, re.DOTALL | re.IGNORECASE)
    if body_match:
        html_content = body_match.group(1)
    
    # Split by tags and process
    lines = re.split(r'(<[^>]+>)', html_content)
    
    current_para = {'type': 'p', 'runs': []}
    current_list = []
    in_list = False
    list_type = None
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Heading tags
        if re.match(r'<h[123]', line, re.IGNORECASE):
            level = int(line[2])
            i += 1
            text = ''
            while i < len(lines) and not re.match(r'</h', lines[i], re.IGNORECASE):
                text += lines[i]
                i += 1
            elements.append({'type': f'h{level}', 'text': strip_tags(text).strip()})
            continue
        
        # Paragraph tag
        elif re.match(r'<p', line, re.IGNORECASE):
            if current_para['runs']:
                elements.append(current_para)
            current_para = {'type': 'p', 'runs': []}
            i += 1
            continue
        
        elif re.match(r'</p>', line, re.IGNORECASE):
            if current_para['runs']:
                elements.append(current_para)
                current_para = {'type': 'p', 'runs': []}
            i += 1
            continue
        
        # List tags
        elif re.match(r'<ul', line, re.IGNORECASE):
            in_list = True
            list_type = 'ul'
            i += 1
            continue
        
        elif re.match(r'<ol', line, re.IGNORECASE):
            in_list = True
            list_type = 'ol'
            i += 1
            continue
        
        elif re.match(r'</[ou]l>', line, re.IGNORECASE):
            if in_list:
                elements.append({'type': list_type, 'items': current_list})
                current_list = []
                in_list = False
            i += 1
            continue
        
        elif re.match(r'<li', line, re.IGNORECASE):
            i += 1
            text = ''
            while i < len(lines) and not re.match(r'</li>', lines[i], re.IGNORECASE):
                text += lines[i]
                i += 1
            current_list.append(strip_tags(text).strip())
            continue
        
        # Text content
        else:
            if not re.match(r'<', line):
                # Plain text
                current_para['runs'].append({'text': line, 'bold': False, 'italic': False, 'underline': False})
        
        i += 1
    
    if current_para['runs']:
        elements.append(current_para)
    
    return elements


def html_to_markdown(html_content):
    """Convert HTML to Markdown"""
    from html2text import html2text
    try:
        markdown_content = html2text(html_content)
        return markdown_content
    except:
        # Fallback: simple regex-based conversion
        text = strip_tags(html_content)
        return text


def html_to_text(html_content):
    """Convert HTML to plain text"""
    return strip_tags(html_content)


def markdown_to_html(markdown_content):
    """Convert Markdown to HTML"""
    return markdown.markdown(markdown_content, extensions=['extra', 'tables'])


def text_to_html(text_content):
    """Convert plain text to HTML"""
    # Escape HTML and wrap in paragraphs
    from django.utils.html import escape
    text_content = escape(text_content)
    # Split by double newlines for paragraphs
    paragraphs = text_content.split('\n\n')
    html = '\n'.join(f'<p>{p.replace(chr(10), "<br>")}</p>' for p in paragraphs if p.strip())
    return html
